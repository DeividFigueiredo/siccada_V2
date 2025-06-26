from flask import Flask, Blueprint, render_template, request, session, redirect, url_for, send_file
import requests
import os
from docx import Document
from datetime import datetime
"""import comtypes.client
import pypandoc
import  pythoncom"""
import psutil
from docx import Document
from reportlab.pdfgen import canvas
import subprocess
from PyPDF2 import PdfMerger
#print(os.environ['PATH'])
import pandas as pd
import re
import shutil

from app.utils import login_required, acesso


docgen_bp= Blueprint('doc_gen', __name__)

# Caminho base para os documentos
pasta_raiz = os.path.join(os.getcwd(), 'static', 'docx_edit', 'dox_template')
pasta_saida = os.path.join(os.getcwd(), 'static', 'docx_edit', 'saida')
caminho_pdf= os.path.join(os.getcwd(), 'static', 'docx_edit', 'saida_pdf', 'documento_final.pdf')

# Garantir que a pasta de saída exista
os.makedirs(pasta_saida, exist_ok=True)

def limpar_pasta_saida(pasta_saida):
    """Remove todos os arquivos da pasta de saída."""
    for arquivo in os.listdir(pasta_saida):
        caminho_arquivo = os.path.join(pasta_saida, arquivo)
        try:
            if os.path.isfile(caminho_arquivo) or os.path.islink(caminho_arquivo):
                os.unlink(caminho_arquivo)  # Remove o arquivo ou link simbólico
                print(f"Arquivo removido: {caminho_arquivo}")
            elif os.path.isdir(caminho_arquivo):
                os.rmdir(caminho_arquivo)  # Remove diretórios vazios
        except Exception as e:
            print(f"Erro ao remover {caminho_arquivo}: {e}")



def encerrar_processos_word():
    """Encerra processos do Microsoft Word que possam estar abertos."""
    for proc in psutil.process_iter(['name']):
        if proc.info['name'] == 'WINWORD.EXE':
            proc.terminate()
            print("Processo do Word encerrado.")

def substituir_placeholders(doc, placeholders):
    for placeholder, valor in placeholders.items():
        for par in doc.paragraphs:
            for run in par.runs:
                if placeholder in run.text:
                    # Substituir o texto mantendo a formatação original
                    run.text = run.text.replace(placeholder, valor)
        
        # Substituir em tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        for run in par.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, valor)
    return doc

def processar_documentos(lista_documentos,plano_nome):

    print(plano_nome)
    pasta_saida_pdf = os.path.join(os.getcwd(), 'static', 'docx_edit', 'saida_pdf')
    pasta_upload= os.path.join(os.getcwd(), 'static', 'uploads')
    caminhos_grupo1 = []  # Para contrato e modelo de reembolso
    caminhos_grupo2 = []  # Para os demais documentos
    limpar_pasta_saida(pasta_saida)
    limpar_pasta_saida(pasta_saida_pdf)
    encerrar_processos_word()
    limpar_pasta_saida(pasta_upload)

     # Adicionar arquivos padrão ao Grupo 2
    manual_orientacao = os.path.join(pasta_raiz, 'manual_orientacao.pdf')
    if os.path.exists(manual_orientacao):
        caminhos_grupo2.append(manual_orientacao)

    

    for doc_info in lista_documentos:
        caminho_modelo = doc_info['caminho']
        nome_saida = doc_info['saida']
        placeholders = doc_info['placeholders']
        
        # Verificar se o arquivo modelo existe
        if not os.path.exists(caminho_modelo):
            print(f"Arquivo modelo não encontrado: {caminho_modelo}")
            continue
        
        # Abrir o documento modelo
        doc = Document(caminho_modelo)
        
        # Substituir os placeholders
        doc = substituir_placeholders(doc, placeholders)
        

        
        # Salvar o documento alterado
        caminho_saida = os.path.join(pasta_saida, nome_saida)
        doc.save(caminho_saida)
        print(f"Documento salvo em: {caminho_saida}")

        caminho_saida_pdf = os.path.join(os.getcwd(), 'static', 'docx_edit', 'saida_pdf', os.path.basename(caminho_saida).replace('.docx', '.pdf'))
        docx_para_pdf(caminho_saida, caminho_saida_pdf)

        # Adicionar o PDF ao grupo correspondente
        if nome_saida in ['contrato.docx', 'modelo_reembolso.docx']:
            caminhos_grupo1.append(caminho_saida_pdf)
        else:
            caminhos_grupo2.append(caminho_saida_pdf)

    # Adicionar o guia de leitura com base no plano
    guia_leitura = os.path.join(pasta_raiz, f'guia_leitura_{plano_nome}.pdf')
    print (plano_nome)
    if os.path.exists(guia_leitura):
        caminhos_grupo2.append(guia_leitura)
    
    # Juntar os PDFs do Grupo 1
    caminho_saida_grupo1 = os.path.join(pasta_saida_pdf, 'grupo1.pdf')
    juntar_pdf(caminhos_grupo1, caminho_saida_grupo1)

    # Juntar os PDFs do Grupo 2
    caminho_saida_grupo2 = os.path.join(pasta_saida_pdf, 'grupo2.pdf')
    juntar_pdf(caminhos_grupo2, caminho_saida_grupo2)

def docx_para_pdf(caminho_docx, caminho_pdf):
    try:
        # Caminho completo para o executável soffice
        soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
        
        # Comando para converter usando o LibreOffice
        comando = [
            soffice_path, '--headless', '--convert-to', 'pdf', '--outdir',
            os.path.dirname(caminho_pdf), caminho_docx
        ]
        # Executar o comando
        subprocess.run(comando, check=True)
        print(f"Documento convertido para PDF: {caminho_pdf}")
    except subprocess.CalledProcessError as e:
        print(f"Erro ao converter para PDF: {e}")
    except FileNotFoundError:
        print("Erro: LibreOffice não encontrado. Certifique-se de que está instalado e no PATH.")

def juntar_pdf(lista_caminhos, caminho_saida):
    merger= PdfMerger()
    try:
        for caminho in lista_caminhos:
            merger.append(caminho)
        merger.write(caminho_saida)
        print(f"PDFs unidos em: {caminho_saida}")
    except Exception as e:
        print(f"Erro ao unir PDFs: {e}")
    finally:
        merger.close()
        

@docgen_bp.route('/gerar_documentos', methods=['GET', 'POST'])
@login_required
@acesso('cadastro','administrativo')
def gerar_documentos():
    if request.method == 'POST':
        # Capturar os dados enviados pelo formulário
        plano_nome = request.form.get('plano_nome', '').strip()
        nome_responsavel = request.form.get('nome_responsavel', '').strip()
        data_hoje = datetime.now().strftime('%Y-%m-%d')
        data_contrato = request.form.get('data_contrato', '').strip()
        numero_contrato = request.form.get('numero_contrato', '').strip()

        session['nome_responsavel_temp'] = nome_responsavel

        dataformatada= datetime.strptime(data_hoje, '%Y-%m-%d').strftime('%d/%m/%Y')
        data_hoje = dataformatada

        datacnt_formatada= datetime.strptime(data_contrato, '%Y-%m-%d').strftime('%d/%m/%Y')
        data_contrato = datacnt_formatada
        

        # Verificar se todos os campos obrigatórios foram preenchidos
        if not all([plano_nome, nome_responsavel, data_contrato, numero_contrato]):
            return "Erro: Todos os campos são obrigatórios.", 400

        # Lista de documentos a serem processados
        lista_documentos = [
            {
                'caminho': os.path.join(pasta_raiz, f'{plano_nome}.docx'),
                'saida': 'contrato.docx',
                'placeholders': {
                    '{{nome_responsavel}}': nome_responsavel,
                    '{{data_hoje}}': data_hoje,
                    '{{data_contrato}}': data_contrato,
                    '{{numero_contrato}}': numero_contrato,
                }
            },
            {
                'caminho': os.path.join(pasta_raiz, 'modelo_reembolso.docx'),
                'saida': 'modelo_reembolso.docx',
                'placeholders': {
                    '{{nome_responsavel}}': nome_responsavel,
                    '{{data_hoje}}': data_hoje,
                    '{{data_contrato}}': data_contrato,
                    '{{numero_contrato}}': numero_contrato,
                }
            },
            {
                'caminho': os.path.join(pasta_raiz, 'protocolo_entrega.docx'),
                'saida': 'protocolo_entrega.docx',
                'placeholders': {
                    '{{nome_responsavel}}': nome_responsavel,
                    '{{data_hoje}}': data_hoje,
                    '{{data_contrato}}': data_contrato,
                    '{{numero_contrato}}': numero_contrato,
                }
            }
        ]

        # Processar os documentos
        processar_documentos(lista_documentos,plano_nome)

        return redirect(url_for('doc_gen.upload_documento'))

    # Caso o método seja GET, renderizar a página normalmente
    return render_template('geradores/gerar_kit.html')

@docgen_bp.route('/upload_documento', methods=['GET', 'POST'])
@login_required
def upload_documento():
    if request.method == 'POST':
        # Verificar se os arquivos foram enviados
        if 'arquivos' not in request.files:
            return "Erro: Nenhum arquivo enviado.", 400

        arquivos = request.files.getlist('arquivos')  # Obter todos os arquivos enviados

        # Caminhos dos PDFs do Grupo 1 e Grupo 2
        caminho_saida_pdf = os.path.join(os.getcwd(), 'static', 'docx_edit', 'saida_pdf')
        caminho_saida_grupo1 = os.path.join(caminho_saida_pdf, 'grupo1.pdf')
        caminho_saida_grupo2 = os.path.join(caminho_saida_pdf, 'grupo2.pdf')

        # Lista para armazenar os caminhos dos PDFs enviados
        caminhos_upload = []

        for arquivo in arquivos:
            # Verificar se o arquivo tem um nome válido
            if arquivo.filename == '':
                return "Erro: Um dos arquivos não foi selecionado.", 400

            # Salvar o arquivo na pasta 'uploads'
            pasta_uploads = os.path.join(os.getcwd(), 'static', 'uploads')
            os.makedirs(pasta_uploads, exist_ok=True)  # Garantir que a pasta exista
            caminho_arquivo = os.path.join(pasta_uploads, arquivo.filename)
            arquivo.save(caminho_arquivo)
            print(f"Arquivo salvo em: {caminho_arquivo}")

            # Converter o arquivo enviado para PDF (se necessário)
            caminho_arquivo_pdf = os.path.join(
                os.getcwd(), 'static', 'docx_edit', 'saida_pdf', arquivo.filename.replace('.docx', '.pdf')
            )
            if arquivo.filename.endswith('.docx'):
                docx_para_pdf(caminho_arquivo, caminho_arquivo_pdf)
            else:
                caminho_arquivo_pdf = caminho_arquivo  # Se já for PDF, usar diretamente

            # Adicionar o caminho do PDF à lista
            caminhos_upload.append(caminho_arquivo_pdf)
        nome_responsavel = session.get('nome_responsavel_temp', 'responsavel')
        # Criar o PDF final unindo Grupo 1, os uploads e Grupo 2
        caminho_pdf_final = os.path.join(caminho_saida_pdf, f'{nome_responsavel}-kit.pdf')
        juntar_pdf([caminho_saida_grupo1] + caminhos_upload + [caminho_saida_grupo2], caminho_pdf_final)

        print(f"Documento final gerado em: {caminho_pdf_final}")

        # Redirecionar para a página de conclusão
        return render_template('geradores/concluido.html')

    # Caso o método seja GET, renderizar a página de upload
    return render_template('geradores/upload_documento.html')

@docgen_bp.route('/download_documento_final', methods=['GET'])
@login_required
def download_documento_final():
    # Recuperar o nome do responsável da sessão
    nome_responsavel = session.get('nome_responsavel_temp', 'responsavel')

    # Caminho do documento final
    caminho_pdf_final = os.path.join(os.getcwd(), 'static', 'docx_edit', 'saida_pdf', f'{nome_responsavel}-kit.pdf')

    # Verificar se o arquivo existe
    if not os.path.exists(caminho_pdf_final):
        return "Erro: Documento final não encontrado.", 404

    # Enviar o arquivo para download
    return send_file(caminho_pdf_final, as_attachment=True)

###################################separação entre rotas para arquivos docx - pdf & xlsx - csv#######################################

def converter_xls_para_csv(caminho_arquivo_excel, pasta_destino):
    """
    Converte um arquivo Excel (.xls ou .xlsx) para CSV.
    """
    try:
        # Ler o arquivo Excel
        df = pd.read_excel(caminho_arquivo_excel)

        # Gerar o caminho do arquivo CSV
        nome_arquivo_csv = os.path.splitext(os.path.basename(caminho_arquivo_excel))[0] + '.csv'
        caminho_arquivo_csv = os.path.join(pasta_destino, nome_arquivo_csv)

        # Salvar como CSV
        df.to_csv(caminho_arquivo_csv, index=False, sep=';', encoding='ISO-8859-1')
        print(f"Arquivo convertido para CSV: {caminho_arquivo_csv}")
        return caminho_arquivo_csv
    except Exception as e:
        print(f"Erro ao converter o arquivo Excel para CSV: {e}")
        return None


@docgen_bp.route("/processar_csv", methods=['POST'])
@login_required
def processar_csv():

    """Processa o arquivo CSV enviado pelo usuário e realiza as transformações necessárias nos números de telefone."""
    # Caminho para upload
    pasta_upload_csv = os.path.join(os.getcwd(), 'static', 'uploads_gen_arch', 'arquivos_originais')
    os.makedirs(pasta_upload_csv, exist_ok=True)  # Garantir que a pasta exista
    limpar_pasta_saida(pasta_upload_csv)

    # Verificar se o arquivo foi enviado
    if 'arquivo_csv' not in request.files or request.files['arquivo_csv'].filename == '':
        print("Erro: Nenhum arquivo enviado.")
        return "Erro: Nenhum arquivo enviado.", 400

    arquivo = request.files['arquivo_csv']
    caminho_arquivo_original = os.path.join(pasta_upload_csv, arquivo.filename)

    try:
        # Salvar o arquivo enviado
        arquivo.save(caminho_arquivo_original)
        print(f"Arquivo salvo em: {caminho_arquivo_original}")

        # Processar o arquivo CSV
        print("Lendo o arquivo CSV...")
        # Forçar a leitura da coluna 'telefone' como string para evitar a conversão em decimal
        df = pd.read_csv(caminho_arquivo_original, sep=';', encoding='ISO-8859-1', dtype={'telefone': str})
        print(f"DataFrame inicial (primeiras 5 linhas):\n{df.head()}")

        # Verificar se a coluna 'telefone' existe
        if 'telefone' not in df.columns:
            print("Erro: A coluna 'telefone' não foi encontrada no arquivo CSV.")
            return "Erro: A coluna 'telefone' não foi encontrada no arquivo CSV.", 400

        # Garantir que a coluna 'telefone' seja tratada como string
        df['telefone'] = df['telefone'].astype(str)

        # Remover todos os caracteres que não sejam números
        df['telefone'] = df['telefone'].str.replace(r'\D', '', regex=True)

        # Função para filtrar números com 8 a 11 dígitos e aplicar as alterações necessárias
        def filtrar_numeros_e_aplicar_prefixo(valor):
            if isinstance(valor, str):
                # Verifica se a string não está vazia ou não contém números
                if not valor.isdigit() or len(valor) == 0:
                    return None
                
                # Se o número tem 8 dígitos, retorna None (para apagar números de 8 dígitos)
                if len(valor) == 8:
                    return None
                
                # Se o número tem 9 dígitos e não começa com '5521', adicionar '5521'
                elif len(valor) == 9 and not valor.startswith('5521'):
                    return '5521' + valor
                
                # Se o número tem 11 dígitos, adicionar '+55'
                elif len(valor) == 11 and not valor.startswith('55'):
                    return '+55' + valor
                
                elif len(valor) > 10 and not valor.startswith('55'):
                    return None
            return None

        # Aplicar a função de filtragem à coluna 'telefone'
        df['telefone'] = df['telefone'].apply(filtrar_numeros_e_aplicar_prefixo)
        
        # Remover as linhas com valores 'None' (números de 8 dígitos ou inválidos)
        df = df.dropna(subset=['telefone'])
        print(f"DataFrame após a filtragem (primeiras 5 linhas):\n{df.head()}")

        # **Novo filtro**: Remover números que começam com '5521' e não têm '9' logo após o prefixo
      

        # Aplicar o filtro adicional à coluna 'telefone'
        
        
        # Remover as linhas com valores 'None' (números que não passam no filtro adicional)
        df = df.dropna(subset=['telefone'])
        print(f"DataFrame após a remoção de números com prefixo '5521' e sem '9' (primeiras 5 linhas):\n{df.head()}")

        # **Adicionar o prefixo 5521 novamente aos números com 9 dígitos**
        

        # Aplicar o filtro para adicionar o prefixo '5521' de volta aos números de 9 dígitos
        
        
        # Verificar se o DataFrame não está vazio antes de salvar
        if not df.empty:
            caminho_filtrado = os.path.join(pasta_upload_csv, 'arquivo_filtrado.csv')
            print(f"Salvando o DataFrame filtrado em: {caminho_filtrado}")
            df.to_csv(caminho_filtrado, index=False)
            print(f"Arquivo filtrado salvo com sucesso em: {caminho_filtrado}")
            return redirect(url_for('doc_gen.confirmar_processamento'))
        else:
            print("Nenhum número de telefone válido encontrado. O DataFrame está vazio.")
            return "Nenhum número de telefone válido encontrado. O DataFrame está vazio.", 400
    except Exception as e:
        print(f"Erro ao processar o arquivo: {e}")
        return f"Erro ao processar o arquivo: {e}", 500

@docgen_bp.route('/dividir_csv', methods=['POST'])  
@login_required
def dividir_csv():
    
    
    import pandas as pd

    
    pasta_upload_csv = os.path.join(os.getcwd(), 'static', 'uploads_gen_arch','arquivos_originais')
    caminho_arquivo_filtrado = os.path.join(pasta_upload_csv, 'arquivo_filtrado.csv')
    pasta_antes = os.path.join(os.getcwd(), 'static', 'uploads_gen_arch')
    pasta_divididos = os.path.join(pasta_antes, 'arquivos_divididos')
    limpar_pasta_saida(pasta_divididos)

    # Verificar se o arquivo filtrado existe
    if not os.path.exists(caminho_arquivo_filtrado):
        return "Erro: Arquivo filtrado não encontrado.", 404

    try:
        # Criar a pasta para os arquivos divididos
        
        
        os.makedirs(pasta_divididos, exist_ok=True)
        limpar_pasta_saida(pasta_divididos)

        # Ler o arquivo filtrado
        df = pd.read_csv(caminho_arquivo_filtrado, encoding='ISO-8859-1', delimiter=',')

        # Número de linhas por arquivo
        linhas_por_arquivo = 80

        # Dividir o DataFrame em partes de 40 linhas
        for i in range(0, len(df), linhas_por_arquivo):
            novo_arquivo = os.path.join(pasta_divididos, f'parte_{i // linhas_por_arquivo + 1}.csv')
            parte_df = df.iloc[i:i + linhas_por_arquivo]
            parte_df.to_csv(novo_arquivo, index=False)
            print(f"Arquivo {novo_arquivo} salvo com {len(parte_df)} linhas.")
    
        # Verificar se os arquivos foram criados
        if not os.listdir(pasta_divididos):
            return "Erro: Nenhum arquivo foi criado na pasta de arquivos divididos.", 500

        return redirect(url_for('doc_gen.baixar_arquivos_divididos'))
    except Exception as e:
        return f"Erro ao dividir o arquivo: {e}"


@docgen_bp.route('/confirmar_processamento', methods=['GET'])
@login_required
def confirmar_processamento():
    return render_template('geradores/confirmar_processamento.html')


@docgen_bp.route('/baixar_arquivos_divididos', methods=['GET'])
@login_required
def baixar_arquivos_divididos():
    # Caminho da pasta onde os arquivos divididos estão
    pasta_upload_csv = os.path.join(os.getcwd(), 'static','uploads_gen_arch')
    pasta_divididos = os.path.join(pasta_upload_csv, 'arquivos_divididos')

    # Verificar se a pasta existe
    if not os.path.exists(pasta_divididos):
        return "Erro: Nenhum arquivo dividido encontrado.", 404

    # Caminho para o arquivo ZIP
    caminho_zip = os.path.join(pasta_upload_csv, 'arquivos_operation','arquivos_divididos.zip')

    # Compactar a pasta em um arquivo ZIP
    try:
        shutil.make_archive(caminho_zip.replace('.zip', ''), 'zip', pasta_divididos)
        print(f"Arquivos compactados em: {caminho_zip}")
    except Exception as e:
        return f"Erro ao compactar os arquivos: {e}"

    # Enviar o arquivo ZIP para download
    return send_file(caminho_zip, as_attachment=True, download_name='faixas_prevencimento.zip')
    
@docgen_bp.route('/processar_csv_form', methods=['GET'])
@login_required
@acesso('cobranca','cadastro')
def processar_csv_form():
    return render_template('geradores/processar_csv.html')

@docgen_bp.route('/confirmar_envio', methods=['GET'])
@login_required
def confirmar_envio():
    return render_template('geradores/confirmar_envio.html')
